"""
Claude-подобный агент на LangChain 1.0
+ Работа с Excel (создание, редактирование, формулы)
+ Память для сохранения контекста
+ ИСПРАВЛЕНА обработка сложных merged cells

v3 — Доработки:
  1. Безопасность: sandbox для bash и python (whitelist, ограничения)
  2. _normalize_merged_cells работает с копией, не портит оригинал
  3. Поиск файлов в нескольких директориях (OUTPUT_DIR, WORK_DIR, абс. путь)
  4. excel_from_csv: авто-определение кодировки и разделителя
  5. extract_multilevel_headers: защита от «рваных» заголовков
  6. DuckDuckGoSearchRun — синглтон, не пересоздаётся
  7. excel_edit_cell: авто-приведение типов (число, формула, строка)
  8. list_files: проверка существования директории
  9. Динамический thread_id для параллельных сессий
  10. Optional[str] вместо str | None для совместимости с Python 3.9+
  11. max_completion_tokens увеличен до 8192
  12. Logging вместо print для предупреждений
"""

import os
import re
import sys
import shutil
import subprocess
import json
import logging
import uuid
from pathlib import Path
from datetime import datetime
from typing import Any, Dict, Optional

from dotenv import load_dotenv

from langchain_openai import ChatOpenAI
from langchain.agents import create_agent
from langchain_core.tools import tool
from langchain_community.tools import DuckDuckGoSearchRun
from langgraph.checkpoint.memory import MemorySaver

# Excel библиотеки
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# PDF библиотеки
try:
    import pymupdf  # PyMuPDF (fitz)
    PDF_AVAILABLE = True
except ImportError:
    try:
        import fitz  # старое имя PyMuPDF
        pymupdf = fitz
        PDF_AVAILABLE = True
    except ImportError:
        PDF_AVAILABLE = False

# Word библиотеки
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Изображения
try:
    from PIL import Image, ImageFilter, ImageEnhance, ImageDraw
    IMAGE_AVAILABLE = True
except ImportError:
    IMAGE_AVAILABLE = False

# Selenium (опционально)
try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.common.keys import Keys
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from selenium.common.exceptions import TimeoutException, NoSuchElementException
    SELENIUM_AVAILABLE = True
except ImportError:
    SELENIUM_AVAILABLE = False

load_dotenv()

# ============ LOGGING ============
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
logger = logging.getLogger("claude_agent")

# ============ CONFIGURATION ============
BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR / "outputs"
OUTPUT_DIR.mkdir(exist_ok=True)

WORK_DIR = BASE_DIR / "work"
WORK_DIR.mkdir(exist_ok=True)

# --- Безопасность: ограничения для bash ---
BASH_BLOCKED_PATTERNS = [
    r"\brm\s+-rf\s+/",       # rm -rf /
    r"\bmkfs\b",             # форматирование дисков
    r"\bdd\s+if=",           # запись в устройства
    r":(){.*};:",            # fork bomb
    r"\bcurl\b.*\|\s*bash",  # curl | bash
    r"\bwget\b.*\|\s*bash",  # wget | bash
    r"\bsudo\b",             # sudo
    r"\bchmod\s+777\s+/",    # chmod 777 /
    r"\b>\s*/etc/",          # перезапись системных файлов
    r"\b>\s*/dev/",          # запись в устройства
]

# --- Безопасность: ограничения для python ---
PYTHON_BLOCKED_IMPORTS = [
    "ctypes", "socket", "http.server", "smtplib",
    "ftplib", "telnetlib", "xmlrpc",
]

SYSTEM_PROMPT = f"""Ты полезный AI-ассистент с доступом к инструментам.

ПОВЕДЕНИЕ:
- Отвечай ТОЛЬКО по существу — давай сразу итоговый ответ
- НЕ показывай ход рассуждений, процесс поиска или промежуточные мысли
- НЕ пиши "Давайте посмотрим...", "Хм, подумаю...", "Я нашёл информацию..." и т.п.
- НЕ выдавай сырые данные из поиска — обработай и структурируй информацию
- Сразу давай конкретный, чёткий и структурированный ответ
- Используй инструменты когда это необходимо

ФОРМАТ ОТВЕТА:
- Отвечай кратко и по делу
- Если нужны цифры — оформи их понятно
- Если нужен список — пронумеруй
- Не копируй результаты поиска дословно — перефразируй и структурируй
- Давай только полезную информацию, убирай мусор и дубли

ИНСТРУМЕНТЫ:
- Для поиска информации используй web_search
- Для чтения содержимого веб-страницы по URL используй fetch_url
- Если пользователь вводит URL — автоматически используй fetch_url чтобы прочитать страницу
- Для работы с Excel используй excel_* инструменты
- Для вычислений используй python_execute
- Будь кратким и естественным
- Помни контекст предыдущих сообщений

EXCEL:
- Если файл содержит многоуровневые заголовки или объединённые ячейки:
  используй excel_read_structured
- Если заголовки плоские:
  используй excel_read
- Для создания сводных таблиц (группировка + агрегация):
  используй excel_create_pivot
- Многоуровневые заголовки объединяй через " | "

PDF:
- Для чтения текста из PDF: pdf_read
- Для информации о PDF: pdf_info
- Для извлечения страниц: pdf_extract_pages

WORD (DOCX):
- Для чтения .docx: docx_read
- Для создания .docx: docx_create (поддержка заголовков # ## ###)

ИЗОБРАЖЕНИЯ:
- Информация: image_info
- Анализ содержимого (Vision): image_analyze — ИСПОЛЬЗУЙ для вопросов "что на картинке"
- Ресайз: image_resize
- Конвертация форматов: image_convert (png, jpg, webp и др.)
- Обрезка: image_crop
- Яркость/контраст/поворот: image_adjust

БРАУЗЕР (если доступен Selenium):
- browser_open: открыть URL в реальном браузере (поддержка JavaScript)
- browser_click: кликнуть по элементу (CSS-селектор, XPath или текст)
- browser_fill: заполнить поле ввода, опционально отправить форму
- browser_extract: извлечь текст из элементов по CSS-селектору
- browser_screenshot: сделать скриншот страницы
- Используй browser_open вместо fetch_url для сайтов с JavaScript
- Используй fetch_url для простых страниц (быстрее)

БЕЗОПАСНОСТЬ:
- bash_execute: запрещены деструктивные команды (rm -rf /, sudo, и т.д.)
- python_execute: запрещены сетевые модули (socket, http.server, и т.д.)
- Файлы ищутся в outputs/ и work/ директориях

ДИРЕКТОРИИ:
- Выходные файлы: {OUTPUT_DIR}
- Рабочая директория: {WORK_DIR}

Текущая дата: {datetime.now().strftime("%B %d, %Y")}"""


# ============ HELPERS: FILE RESOLUTION ============

def _resolve_file(filename: str, must_exist: bool = True) -> Optional[Path]:
    """
    Ищет файл в нескольких директориях:
    1. OUTPUT_DIR
    2. WORK_DIR
    3. Абсолютный путь (если указан)

    Возвращает Path или None если файл не найден.
    """
    candidates = [
        OUTPUT_DIR / Path(filename).name,
        WORK_DIR / Path(filename).name,
    ]

    # Если передан абсолютный путь — тоже проверяем
    abs_path = Path(filename)
    if abs_path.is_absolute():
        candidates.insert(0, abs_path)

    for p in candidates:
        if p.exists():
            return p

    if not must_exist:
        # Для создания — возвращаем путь в OUTPUT_DIR
        return OUTPUT_DIR / Path(filename).name

    return None


# ============ HELPERS: SECURITY ============

def _check_bash_safety(command: str) -> Optional[str]:
    """Проверяет команду bash на опасные паттерны. Возвращает ошибку или None."""
    for pattern in BASH_BLOCKED_PATTERNS:
        if re.search(pattern, command, re.IGNORECASE):
            return f"⛔ Команда заблокирована (опасный паттерн: {pattern})"
    return None


def _check_python_safety(code: str) -> Optional[str]:
    """Проверяет Python-код на запрещённые импорты. Возвращает ошибку или None."""
    # Блокируем __import__() для обхода проверок
    if "__import__" in code:
        return "⛔ Запрещено: __import__()"
    if "importlib" in code:
        return "⛔ Запрещено: importlib"
    for mod in PYTHON_BLOCKED_IMPORTS:
        if re.search(rf"\bimport\s+{re.escape(mod)}\b", code) or \
           re.search(rf"\bfrom\s+{re.escape(mod)}\b", code):
            return f"⛔ Запрещённый модуль: {mod}"
    return None


# ============ EXCEL TOOLS ============

def _normalize_merged_cells(filepath: Path) -> Path:
    """
    Безопасно убирает merged cells, работая с КОПИЕЙ файла.
    Возвращает путь к нормализованной копии (или оригинал, если merged cells нет).
    """
    try:
        wb = openpyxl.load_workbook(filepath)
        ws = wb.active

        merged_list = list(ws.merged_cells.ranges)

        if not merged_list:
            wb.close()
            return filepath  # Нет merged cells — возвращаем оригинал

        # Есть merged cells — работаем с копией
        tmp_path = filepath.with_suffix(".normalized.xlsx")
        shutil.copy2(filepath, tmp_path)

        wb_tmp = openpyxl.load_workbook(tmp_path)
        ws_tmp = wb_tmp.active

        merged_list_tmp = list(ws_tmp.merged_cells.ranges)

        # Собираем данные
        cells_data = []
        for merged in merged_list_tmp:
            min_row, min_col, max_row, max_col = merged.bounds
            value = ws_tmp.cell(row=min_row, column=min_col).value
            cells_data.append((merged, value, min_row, min_col, max_row, max_col))

        # Разъединяем
        for merged, _, _, _, _, _ in cells_data:
            try:
                ws_tmp.unmerge_cells(str(merged))
            except Exception:
                pass

        # Заполняем значениями
        for _, value, min_row, min_col, max_row, max_col in cells_data:
            for r in range(min_row, max_row + 1):
                for c in range(min_col, max_col + 1):
                    ws_tmp.cell(row=r, column=c, value=value)

        wb_tmp.save(tmp_path)
        wb.close()
        wb_tmp.close()
        return tmp_path

    except Exception as e:
        logger.warning(f"Предупреждение при нормализации: {e}")
        return filepath  # Возвращаем оригинал если что-то пошло не так


def _cleanup_normalized(tmp_path: Path, original_path: Path) -> None:
    """Удаляет временный нормализованный файл, если он отличается от оригинала."""
    if tmp_path != original_path and tmp_path.exists():
        try:
            tmp_path.unlink()
        except Exception:
            pass


@tool
def excel_create(filename: str, data: str, sheet_name: str = "Sheet1") -> str:
    """Создать Excel файл с данными.

    Args:
        filename: Имя файла (будет создан в outputs/)
        data: JSON-массив строк, например: [["Имя", "Возраст"], ["Анна", 25]]
        sheet_name: Имя листа (по умолчанию Sheet1)
    """
    if not EXCEL_AVAILABLE:
        return "Ошибка: openpyxl не установлен"

    try:
        rows = json.loads(data)
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        for row_idx, row_data in enumerate(rows, 1):
            for col_idx, value in enumerate(row_data, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

        # Автоширина колонок
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            ws.column_dimensions[column_letter].width = min(max_length + 2, 50)

        filepath = OUTPUT_DIR / filename
        wb.save(filepath)
        return f"✓ Excel создан: {filepath} ({len(rows)} строк)"
    except json.JSONDecodeError as e:
        return f"Ошибка парсинга JSON data: {e}"
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def excel_add_formulas(filename: str, formulas: str) -> str:
    """Добавить формулы в Excel.

    Args:
        filename: Имя файла
        formulas: JSON-массив, например: [{"cell": "C2", "formula": "=A2+B2"}]
    """
    if not EXCEL_AVAILABLE:
        return "Ошибка: openpyxl не установлен"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        wb = openpyxl.load_workbook(filepath)
        ws = wb.active

        formula_list = json.loads(formulas)
        for item in formula_list:
            ws[item["cell"]] = item["formula"]

        wb.save(filepath)
        return f"✓ Добавлено {len(formula_list)} формул в {filepath.name}"
    except json.JSONDecodeError as e:
        return f"Ошибка парсинга JSON formulas: {e}"
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def excel_style(filename: str, styles: str) -> str:
    """Применить стили к Excel.

    Args:
        filename: Имя файла
        styles: JSON со стилями. Поддерживаемые ключи:
            header_row (int), header_color (hex), header_font_color (hex),
            freeze_panes (str, напр. "A2"), borders (bool)
    """
    if not EXCEL_AVAILABLE:
        return "Ошибка: openpyxl не установлен"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        wb = openpyxl.load_workbook(filepath)
        ws = wb.active

        style_dict = json.loads(styles)

        if "header_row" in style_dict:
            header_row = style_dict["header_row"]
            bg_color = style_dict.get("header_color", "4472C4")
            font_color = style_dict.get("header_font_color", "FFFFFF")

            for cell in ws[header_row]:
                cell.font = Font(bold=True, color=font_color)
                cell.fill = PatternFill(start_color=bg_color, end_color=bg_color, fill_type="solid")
                cell.alignment = Alignment(horizontal="center", vertical="center")

        if "freeze_panes" in style_dict:
            ws.freeze_panes = style_dict["freeze_panes"]

        if style_dict.get("borders", False):
            thin_border = Border(
                left=Side(style='thin'), right=Side(style='thin'),
                top=Side(style='thin'), bottom=Side(style='thin')
            )
            for row in ws.iter_rows():
                for cell in row:
                    cell.border = thin_border

        wb.save(filepath)
        return f"✓ Стили применены к {filepath.name}"
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def excel_read(filename: str, sheet_name: str = None) -> str:
    """Чтение Excel с автоматической обработкой merged cells.

    Args:
        filename: Имя файла (ищет в outputs/, work/ и по абсолютному пути)
        sheet_name: Имя листа (опционально, по умолчанию — первый)
    """
    if not EXCEL_AVAILABLE:
        return "Ошибка: openpyxl / pandas не установлен"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename} (проверены: outputs/, work/)"

        # Нормализуем merged cells на КОПИИ
        normalized = _normalize_merged_cells(filepath)

        try:
            df = pd.read_excel(normalized, sheet_name=sheet_name or 0, dtype=str)
            df = df.dropna(how="all").dropna(axis=1, how="all")

            preview = df.head(20).to_string(index=False)
            return (
                f"Файл: {filepath.name}\n"
                f"Строк: {len(df)}\n"
                f"Колонок: {len(df.columns)}\n\n"
                f"Первые строки:\n{preview}"
            )
        finally:
            _cleanup_normalized(normalized, filepath)

    except Exception as e:
        return f"Ошибка чтения: {e}"


@tool
def excel_edit_cell(filename: str, cell: str, value: str) -> str:
    """Изменить ячейку в Excel.

    Автоматически определяет тип: число, формула или строка.

    Args:
        filename: Имя файла
        cell: Адрес ячейки (напр. "A1", "B5")
        value: Значение (строка, число или формула начинающаяся с =)
    """
    if not EXCEL_AVAILABLE:
        return "Ошибка: openpyxl не установлен"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        wb = openpyxl.load_workbook(filepath)
        ws = wb.active

        # Авто-приведение типов
        parsed_value: Any = value
        if isinstance(value, str):
            if value.startswith("="):
                # Формула — оставляем как есть
                parsed_value = value
            else:
                # Пробуем число
                try:
                    if "." in value:
                        parsed_value = float(value)
                    else:
                        parsed_value = int(value)
                except ValueError:
                    parsed_value = value  # Строка

        ws[cell] = parsed_value
        wb.save(filepath)

        type_label = "формула" if str(value).startswith("=") else type(parsed_value).__name__
        return f"✓ Ячейка {cell} = {parsed_value} ({type_label})"
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def excel_from_csv(csv_filename: str, excel_filename: str) -> str:
    """Конвертировать CSV в Excel.

    Автоматически определяет кодировку (utf-8, cp1251, latin-1)
    и разделитель (запятая, точка с запятой, табуляция).

    Args:
        csv_filename: Имя CSV-файла
        excel_filename: Имя выходного Excel-файла
    """
    if not EXCEL_AVAILABLE:
        return "Ошибка: pandas не установлен"

    try:
        csv_path = _resolve_file(csv_filename)
        if not csv_path:
            return f"CSV не найден: {csv_filename}"

        # Определяем кодировку
        df = None
        encodings = ["utf-8", "cp1251", "latin-1"]
        separators = [",", ";", "\t"]

        for enc in encodings:
            for sep in separators:
                try:
                    df = pd.read_csv(csv_path, encoding=enc, sep=sep)
                    # Проверяем что прочиталось разумно (> 1 колонки или 1 строка)
                    if len(df.columns) > 1 or len(df) > 0:
                        break
                except (UnicodeDecodeError, pd.errors.ParserError):
                    continue
            if df is not None and len(df.columns) > 1:
                break

        if df is None:
            return "Не удалось прочитать CSV — неизвестная кодировка или формат"

        excel_path = OUTPUT_DIR / excel_filename

        with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Data')
            ws = writer.sheets['Data']
            for cell in ws[1]:
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color="4472C4", fill_type="solid")

        return (
            f"✓ Создан {excel_filename}\n"
            f"  Строк: {len(df)}, Колонок: {len(df.columns)}\n"
            f"  Кодировка: {enc}, Разделитель: {repr(sep)}"
        )
    except Exception as e:
        return f"Ошибка: {e}"


def extract_multilevel_headers(ws, max_header_rows: int = 3) -> list:
    """Извлекает многоуровневые заголовки.

    Защита от «рваных» заголовков (разное количество колонок в строках).
    """
    header_matrix = []

    for row in ws.iter_rows(min_row=1, max_row=max_header_rows, values_only=True):
        header_matrix.append([str(c).strip() if c else "" for c in row])

    if not header_matrix:
        return []

    # Определяем максимальную ширину по всем строкам
    max_cols = max(len(row) for row in header_matrix)

    # Выравниваем строки до одинаковой длины
    for i, row in enumerate(header_matrix):
        if len(row) < max_cols:
            header_matrix[i] = row + [""] * (max_cols - len(row))

    headers = []
    for col_idx in range(max_cols):
        parts = []
        for row_idx in range(len(header_matrix)):
            val = header_matrix[row_idx][col_idx]
            if val and (not parts or val != parts[-1]):
                parts.append(val)
        headers.append(" | ".join(parts) if parts else f"Column_{col_idx + 1}")

    return headers


@tool
def excel_read_structured(filename: str, header_rows: int = 2) -> str:
    """Чтение Excel с многоуровневыми заголовками.

    Args:
        filename: Имя файла
        header_rows: Количество строк заголовков (по умолчанию 2)
    """
    if not EXCEL_AVAILABLE:
        return "Ошибка: openpyxl / pandas не установлен"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        normalized = _normalize_merged_cells(filepath)

        try:
            wb = openpyxl.load_workbook(normalized)
            ws = wb.active

            headers = extract_multilevel_headers(ws, max_header_rows=header_rows)
            wb.close()

            df = pd.read_excel(normalized, skiprows=header_rows, header=None, dtype=str)
            df.columns = headers[:len(df.columns)]
            df = df.dropna(how="all")

            preview = df.head(20).to_string(index=False)

            return (
                f"Файл: {filepath.name}\n"
                f"Заголовков уровней: {header_rows}\n"
                f"Колонок: {len(df.columns)}\n\n"
                f"Имена колонок:\n" + "\n".join(f"- {c}" for c in df.columns)
                + "\n\nПервые строки:\n" + preview
            )
        finally:
            _cleanup_normalized(normalized, filepath)

    except Exception as e:
        return f"Ошибка: {e}"


# ============ GENERAL TOOLS ============

# Синглтон для DuckDuckGo — не пересоздаётся при каждом вызове
_search_instance: Optional[DuckDuckGoSearchRun] = None


def _get_search() -> DuckDuckGoSearchRun:
    global _search_instance
    if _search_instance is None:
        _search_instance = DuckDuckGoSearchRun()
    return _search_instance


@tool
def web_search(query: str) -> str:
    """Поиск в интернете через DuckDuckGo.

    Args:
        query: Поисковый запрос
    """
    try:
        return _get_search().run(query)
    except Exception as e:
        return f"Ошибка поиска: {e}"


@tool
def bash_execute(command: str) -> str:
    """Выполнить bash-команду в рабочей директории.

    Заблокированы деструктивные команды (rm -rf /, sudo, fork bomb и т.д.).

    Args:
        command: Команда для выполнения
    """
    # Проверка безопасности
    safety_error = _check_bash_safety(command)
    if safety_error:
        return safety_error

    try:
        result = subprocess.run(
            command, shell=True, capture_output=True, text=True,
            timeout=30, cwd=str(WORK_DIR)
        )
        parts = []
        if result.stdout:
            parts.append(f"STDOUT:\n{result.stdout[:5000]}")  # Лимит вывода
        if result.stderr:
            parts.append(f"STDERR:\n{result.stderr[:2000]}")
        if result.returncode != 0:
            parts.append(f"Exit code: {result.returncode}")
        return "\n".join(parts) or "Выполнено (нет вывода)"
    except subprocess.TimeoutExpired:
        return "⛔ Команда превысила таймаут (30 сек)"
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def create_file(path: str, content: str) -> str:
    """Создать текстовый файл в outputs/.

    Args:
        path: Имя файла (будет создан в outputs/)
        content: Содержимое файла
    """
    try:
        # Безопасность: только имя файла, без path traversal
        safe_name = Path(path).name
        if not safe_name:
            return "Ошибка: некорректное имя файла"

        file_path = OUTPUT_DIR / safe_name
        file_path.write_text(content, encoding="utf-8")
        return f"✓ Файл создан: {file_path}"
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def view_file(path: str) -> str:
    """Прочитать текстовый файл.

    Ищет в outputs/, work/ и по абсолютному пути.

    Args:
        path: Путь или имя файла
    """
    filepath = _resolve_file(path)
    if not filepath:
        return f"Файл не найден: {path} (проверены: outputs/, work/)"

    try:
        content = filepath.read_text(encoding="utf-8")
        # Ограничиваем вывод
        if len(content) > 10_000:
            content = content[:10_000] + f"\n\n... (обрезано, всего {len(content)} символов)"
        return f"[{filepath}]\n\n{content}"
    except UnicodeDecodeError:
        return f"Файл {filepath} — бинарный, невозможно прочитать как текст"
    except Exception as e:
        return f"Ошибка чтения: {e}"


@tool
def list_files(directory: str = "outputs") -> str:
    """Список файлов в директории.

    Args:
        directory: "outputs" или "work"
    """
    target = OUTPUT_DIR if directory == "outputs" else WORK_DIR

    if not target.exists():
        return f"Директория не существует: {target}"

    try:
        items = sorted(target.iterdir())
    except PermissionError:
        return f"Нет доступа к: {target}"

    if not items:
        return f"[{target}]\nПусто"

    lines = []
    for item in items:
        if item.name.startswith("."):
            continue  # Скрываем dotfiles
        if item.is_file():
            size = item.stat().st_size
            # Человекочитаемый размер
            if size < 1024:
                size_str = f"{size} B"
            elif size < 1024 * 1024:
                size_str = f"{size / 1024:.1f} KB"
            else:
                size_str = f"{size / (1024 * 1024):.1f} MB"
            lines.append(f"FILE  {size_str:>10}  {item.name}")
        else:
            lines.append(f"DIR   {'':>10}  {item.name}/")

    return f"[{target}]\n" + "\n".join(lines)


@tool
def python_execute(code: str) -> str:
    """Выполнить Python-код в изолированной среде.

    Заблокированы сетевые модули (socket, http.server, smtplib и т.д.).

    Args:
        code: Python-код для выполнения
    """
    # Проверка безопасности
    safety_error = _check_python_safety(code)
    if safety_error:
        return safety_error

    try:
        script = WORK_DIR / f"_exec_{datetime.now().timestamp()}.py"
        script.write_text(code, encoding="utf-8")

        result = subprocess.run(
            [sys.executable, str(script)],
            capture_output=True, text=True, timeout=30, cwd=str(WORK_DIR)
        )

        script.unlink(missing_ok=True)

        parts = []
        if result.stdout:
            parts.append(f"Вывод:\n{result.stdout[:5000]}")
        if result.stderr:
            parts.append(f"Ошибки:\n{result.stderr[:2000]}")
        if result.returncode != 0:
            parts.append(f"Exit code: {result.returncode}")
        return "\n".join(parts) or "Выполнено (нет вывода)"

    except subprocess.TimeoutExpired:
        script.unlink(missing_ok=True)
        return "⛔ Скрипт превысил таймаут (30 сек)"
    except Exception as e:
        return f"Ошибка: {e}"


# ============ СВОДНЫЕ ТАБЛИЦЫ (PIVOT) ============

@tool
def excel_create_pivot(
    source_file: str,
    output_file: str,
    row_fields: str,
    column_fields: str = None,
    value_field: str = None,
    agg_func: str = "sum",
    show_totals: bool = True
) -> str:
    """
    Создать сводную таблицу (pivot table) из данных Excel.

    Args:
        source_file: Имя исходного файла
        output_file: Имя файла для сохранения сводной
        row_fields: JSON список полей для строк: ["Категория", "Продукт"]
        column_fields: JSON список полей для столбцов: ["Месяц"] (опционально)
        value_field: Поле для агрегации ("Сумма", "Количество")
        agg_func: Функция: "sum", "mean", "count", "min", "max"
        show_totals: Показывать итоги
    """
    if not EXCEL_AVAILABLE:
        return "Ошибка: pandas не установлен"

    # Валидация agg_func
    allowed_funcs = {"sum", "mean", "count", "min", "max", "median", "std"}
    if agg_func not in allowed_funcs:
        return f"Ошибка: agg_func должна быть одной из: {', '.join(allowed_funcs)}"

    try:
        source_path = _resolve_file(source_file)
        if not source_path:
            return f"Файл не найден: {source_file}"

        df = pd.read_excel(source_path)

        # Парсим поля
        rows = json.loads(row_fields) if isinstance(row_fields, str) else row_fields
        cols = json.loads(column_fields) if column_fields and isinstance(column_fields, str) else column_fields

        # Валидация полей
        missing_fields = [f for f in rows if f not in df.columns]
        if value_field and value_field not in df.columns:
            missing_fields.append(value_field)
        if cols:
            missing_fields.extend(f for f in cols if f not in df.columns)
        if missing_fields:
            return (
                f"Поля не найдены в данных: {', '.join(missing_fields)}\n"
                f"Доступные колонки: {', '.join(df.columns)}"
            )

        # Создаём сводную
        pivot_params = {
            'values': value_field,
            'index': rows,
            'aggfunc': agg_func,
        }

        if cols:
            pivot_params['columns'] = cols

        if show_totals:
            pivot_params['margins'] = True
            pivot_params['margins_name'] = 'Итого'

        pivot = df.pivot_table(**pivot_params)

        # Сохраняем с форматированием
        output_path = OUTPUT_DIR / output_file

        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            pivot.to_excel(writer, sheet_name='Сводная')
            ws = writer.sheets['Сводная']

            # Форматирование заголовков
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")

            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center")

            for row in ws.iter_rows(min_row=2, max_row=ws.max_row, min_col=1, max_col=1):
                for cell in row:
                    cell.font = Font(bold=True)

            # Форматирование итогов
            if show_totals:
                last_row = ws.max_row
                total_fill = PatternFill(start_color="FFD966", end_color="FFD966", fill_type="solid")
                for cell in ws[last_row]:
                    cell.fill = total_fill
                    cell.font = Font(bold=True)

            # Границы
            thin_border = Border(
                left=Side(style='thin'), right=Side(style='thin'),
                top=Side(style='thin'), bottom=Side(style='thin')
            )

            for row in ws.iter_rows(min_row=1, max_row=ws.max_row):
                for cell in row:
                    cell.border = thin_border

            # Автоширина
            for column in ws.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                ws.column_dimensions[column_letter].width = min(max_length + 2, 40)

        total_rows = len(pivot)
        total_cols = len(pivot.columns) if hasattr(pivot, 'columns') else 1

        return (
            f"✓ Сводная таблица: {output_file}\n\n"
            f"Строки: {', '.join(rows)}\n"
            f"Столбцы: {', '.join(cols) if cols else 'нет'}\n"
            f"Агрегация: {agg_func}({value_field})\n"
            f"Размер: {total_rows} × {total_cols}\n"
            f"Итоги: {'да' if show_totals else 'нет'}"
        )

    except json.JSONDecodeError as e:
        return f"Ошибка парсинга JSON полей: {e}"
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def excel_pivot_analyze(source_file: str) -> str:
    """
    Проанализировать данные и предложить варианты сводных таблиц.

    Args:
        source_file: Имя файла с данными
    """
    if not EXCEL_AVAILABLE:
        return "Ошибка: pandas не установлен"

    try:
        source_path = _resolve_file(source_file)
        if not source_path:
            return f"Файл не найден: {source_file}"

        df = pd.read_excel(source_path)

        analysis = []
        analysis.append(f"Файл: {source_path.name}")
        analysis.append(f"Строк: {len(df)}, Колонок: {len(df.columns)}\n")

        analysis.append("Поля:")
        for col in df.columns:
            dtype = df[col].dtype
            unique = df[col].nunique()
            nulls = df[col].isnull().sum()

            if unique < 50 and dtype == 'object':
                tip = "→ для группировки (строки/столбцы)"
            elif dtype in ['int64', 'float64']:
                tip = "→ для агрегации (значения)"
            else:
                tip = ""

            analysis.append(
                f"  • {col}: тип={dtype}, уникальных={unique}, пустых={nulls} {tip}"
            )

        # Рекомендации
        categorical = [c for c in df.columns if df[c].dtype == 'object' and df[c].nunique() < 50]
        numeric = [c for c in df.columns if df[c].dtype in ['int64', 'float64']]

        analysis.append("\nРекомендации:")
        if categorical and numeric:
            analysis.append(f"  • Группировка: {', '.join(categorical[:3])}")
            analysis.append(f"  • Агрегация: {', '.join(numeric[:3])}")

            # Пример команды
            analysis.append(f"\n  Пример вызова:")
            analysis.append(
                f'  excel_create_pivot("{source_path.name}", "pivot.xlsx", '
                f"'{json.dumps(categorical[:1])}', None, \"{numeric[0]}\", \"sum\")"
            )
        else:
            analysis.append("  • Недостаточно данных для сводной таблицы")

        return "\n".join(analysis)

    except Exception as e:
        return f"Ошибка: {e}"


# ============ PDF TOOLS ============

@tool
def pdf_read(filename: str, max_pages: int = 50) -> str:
    """Прочитать текст из PDF-файла.

    Args:
        filename: Имя PDF-файла
        max_pages: Максимум страниц для чтения (по умолчанию 50)
    """
    if not PDF_AVAILABLE:
        return "Ошибка: PyMuPDF не установлен. pip install pymupdf"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        doc = pymupdf.open(str(filepath))
        total = len(doc)
        pages_to_read = min(total, max_pages)

        text_parts = []
        for i in range(pages_to_read):
            page = doc[i]
            text = page.get_text()
            if text.strip():
                text_parts.append(f"--- Страница {i + 1} ---\n{text.strip()}")

        doc.close()

        if not text_parts:
            return f"PDF {filepath.name}: {total} страниц, но текст не извлечён (возможно, скан)"

        content = "\n\n".join(text_parts)
        # Ограничиваем вывод
        if len(content) > 15000:
            content = content[:15000] + f"\n\n... (обрезано, прочитано {pages_to_read}/{total} стр.)"

        return (
            f"PDF: {filepath.name}\n"
            f"Страниц: {total} (прочитано: {pages_to_read})\n\n"
            f"{content}"
        )
    except Exception as e:
        return f"Ошибка чтения PDF: {e}"


@tool
def pdf_info(filename: str) -> str:
    """Получить информацию о PDF: кол-во страниц, метаданные, размер.

    Args:
        filename: Имя PDF-файла
    """
    if not PDF_AVAILABLE:
        return "Ошибка: PyMuPDF не установлен"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        doc = pymupdf.open(str(filepath))
        meta = doc.metadata

        size_mb = filepath.stat().st_size / (1024 * 1024)

        info = [
            f"Файл: {filepath.name}",
            f"Размер: {size_mb:.2f} MB",
            f"Страниц: {len(doc)}",
        ]

        if meta:
            for key in ["title", "author", "subject", "creator"]:
                val = meta.get(key, "")
                if val:
                    info.append(f"{key.capitalize()}: {val}")

        doc.close()
        return "\n".join(info)
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def pdf_extract_pages(filename: str, pages: str, output_filename: str) -> str:
    """Извлечь определённые страницы из PDF в новый файл.

    Args:
        filename: Исходный PDF
        pages: Номера страниц, напр. "1,3,5-10"
        output_filename: Имя выходного PDF
    """
    if not PDF_AVAILABLE:
        return "Ошибка: PyMuPDF не установлен"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        # Парсим номера страниц
        page_nums = []
        for part in pages.split(","):
            part = part.strip()
            if "-" in part:
                start, end = part.split("-", 1)
                page_nums.extend(range(int(start) - 1, int(end)))
            else:
                page_nums.append(int(part) - 1)

        doc = pymupdf.open(str(filepath))
        new_doc = pymupdf.open()

        for pn in page_nums:
            if 0 <= pn < len(doc):
                new_doc.insert_pdf(doc, from_page=pn, to_page=pn)

        output_path = OUTPUT_DIR / output_filename
        new_doc.save(str(output_path))
        new_doc.close()
        doc.close()

        return f"✓ Извлечено {len(page_nums)} страниц → {output_filename}"
    except Exception as e:
        return f"Ошибка: {e}"


# ============ WORD (DOCX) TOOLS ============

@tool
def docx_read(filename: str) -> str:
    """Прочитать текст из Word-документа (.docx).

    Args:
        filename: Имя .docx файла
    """
    if not DOCX_AVAILABLE:
        return "Ошибка: python-docx не установлен. pip install python-docx"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        doc = DocxDocument(str(filepath))

        parts = []

        # Параграфы
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style = para.style.name if para.style else ""
                if "Heading" in style:
                    level = style.replace("Heading ", "").replace("Heading", "1")
                    parts.append(f"{'#' * int(level)} {text}")
                else:
                    parts.append(text)

        # Таблицы
        for i, table in enumerate(doc.tables):
            parts.append(f"\n--- Таблица {i + 1} ---")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(" | ".join(cells))

        content = "\n\n".join(parts)
        if len(content) > 15000:
            content = content[:15000] + "\n\n... (обрезано)"

        return (
            f"Word: {filepath.name}\n"
            f"Параграфов: {len(doc.paragraphs)}\n"
            f"Таблиц: {len(doc.tables)}\n\n"
            f"{content}"
        )
    except Exception as e:
        return f"Ошибка чтения DOCX: {e}"


@tool
def docx_create(filename: str, content: str, title: str = "") -> str:
    """Создать Word-документ (.docx).

    Args:
        filename: Имя файла (напр. "report.docx")
        content: Текст документа. Форматирование:
            # Заголовок 1
            ## Заголовок 2
            ### Заголовок 3
            Обычный текст — параграф
            --- — разделитель страниц
        title: Заголовок документа (опционально)
    """
    if not DOCX_AVAILABLE:
        return "Ошибка: python-docx не установлен"

    try:
        doc = DocxDocument()

        if title:
            doc.add_heading(title, level=0)

        for line in content.split("\n"):
            line_stripped = line.strip()

            if not line_stripped:
                continue
            elif line_stripped == "---":
                doc.add_page_break()
            elif line_stripped.startswith("### "):
                doc.add_heading(line_stripped[4:], level=3)
            elif line_stripped.startswith("## "):
                doc.add_heading(line_stripped[3:], level=2)
            elif line_stripped.startswith("# "):
                doc.add_heading(line_stripped[2:], level=1)
            else:
                doc.add_paragraph(line_stripped)

        output_path = OUTPUT_DIR / filename
        doc.save(str(output_path))

        return f"✓ Word создан: {filename} ({len(doc.paragraphs)} параграфов)"
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def docx_to_pdf(filename: str, output_filename: str = "") -> str:
    """Информация о конвертации DOCX в PDF.

    Args:
        filename: Имя .docx файла
        output_filename: Имя выходного PDF (опционально)
    """
    return (
        "Прямая конвертация DOCX→PDF требует LibreOffice.\n"
        "Если LibreOffice установлен, используйте bash_execute:\n"
        f'  bash_execute("libreoffice --headless --convert-to pdf {filename}")\n\n'
        "Или создайте PDF через python_execute с библиотекой reportlab/fpdf."
    )


# ============ IMAGE TOOLS ============

@tool
def image_info(filename: str) -> str:
    """Получить информацию об изображении: размер, формат, режим.

    Args:
        filename: Имя файла изображения
    """
    if not IMAGE_AVAILABLE:
        return "Ошибка: Pillow не установлен. pip install Pillow"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        img = Image.open(filepath)
        size_kb = filepath.stat().st_size / 1024

        info = [
            f"Файл: {filepath.name}",
            f"Формат: {img.format or filepath.suffix}",
            f"Размер: {img.width} × {img.height} px",
            f"Режим: {img.mode}",
            f"Вес: {size_kb:.1f} KB",
        ]

        if hasattr(img, 'info'):
            if 'dpi' in img.info:
                info.append(f"DPI: {img.info['dpi']}")

        img.close()
        return "\n".join(info)
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def image_resize(filename: str, width: int, height: int = 0, output_filename: str = "") -> str:
    """Изменить размер изображения.

    Args:
        filename: Исходный файл
        width: Новая ширина в пикселях
        height: Новая высота (0 = пропорционально)
        output_filename: Имя выходного файла (пустое = перезаписать)
    """
    if not IMAGE_AVAILABLE:
        return "Ошибка: Pillow не установлен"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        img = Image.open(filepath)

        if height == 0:
            ratio = width / img.width
            height = int(img.height * ratio)

        resized = img.resize((width, height), Image.LANCZOS)

        out_name = output_filename or filepath.name
        out_path = OUTPUT_DIR / out_name
        resized.save(out_path)

        img.close()
        resized.close()

        return f"✓ {out_name}: {width}×{height} px"
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def image_convert(filename: str, output_format: str, output_filename: str = "") -> str:
    """Конвертировать изображение в другой формат.

    Args:
        filename: Исходный файл
        output_format: Целевой формат: png, jpg, jpeg, webp, bmp, tiff
        output_filename: Имя выходного файла (опционально)
    """
    if not IMAGE_AVAILABLE:
        return "Ошибка: Pillow не установлен"

    allowed = {"png", "jpg", "jpeg", "webp", "bmp", "tiff", "gif"}
    fmt = output_format.lower().strip(".")
    if fmt not in allowed:
        return f"Неподдерживаемый формат. Доступные: {', '.join(allowed)}"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        img = Image.open(filepath)

        # RGBA → RGB для jpg
        if fmt in ("jpg", "jpeg") and img.mode == "RGBA":
            bg = Image.new("RGB", img.size, (255, 255, 255))
            bg.paste(img, mask=img.split()[3])
            img = bg

        if not output_filename:
            output_filename = filepath.stem + "." + fmt

        out_path = OUTPUT_DIR / output_filename
        img.save(out_path)
        img.close()

        size_kb = out_path.stat().st_size / 1024
        return f"✓ {output_filename} ({size_kb:.1f} KB)"
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def image_crop(filename: str, left: int, top: int, right: int, bottom: int,
               output_filename: str = "") -> str:
    """Обрезать изображение.

    Args:
        filename: Исходный файл
        left: Левая граница (px)
        top: Верхняя граница (px)
        right: Правая граница (px)
        bottom: Нижняя граница (px)
        output_filename: Имя выходного файла
    """
    if not IMAGE_AVAILABLE:
        return "Ошибка: Pillow не установлен"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        img = Image.open(filepath)
        cropped = img.crop((left, top, right, bottom))

        out_name = output_filename or f"cropped_{filepath.name}"
        out_path = OUTPUT_DIR / out_name
        cropped.save(out_path)

        img.close()
        cropped.close()

        return f"✓ {out_name}: {right - left}×{bottom - top} px"
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def image_adjust(filename: str, brightness: float = 1.0, contrast: float = 1.0,
                 sharpness: float = 1.0, rotate: int = 0,
                 output_filename: str = "") -> str:
    """Настроить яркость, контраст, резкость или повернуть изображение.

    Args:
        filename: Исходный файл
        brightness: Яркость (1.0 = без изменений, 1.5 = ярче, 0.5 = темнее)
        contrast: Контраст (1.0 = без изменений)
        sharpness: Резкость (1.0 = без изменений, 2.0 = резче)
        rotate: Угол поворота в градусах (90, 180, 270 или произвольный)
        output_filename: Имя выходного файла
    """
    if not IMAGE_AVAILABLE:
        return "Ошибка: Pillow не установлен"

    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        img = Image.open(filepath)

        if brightness != 1.0:
            img = ImageEnhance.Brightness(img).enhance(brightness)
        if contrast != 1.0:
            img = ImageEnhance.Contrast(img).enhance(contrast)
        if sharpness != 1.0:
            img = ImageEnhance.Sharpness(img).enhance(sharpness)
        if rotate:
            img = img.rotate(rotate, expand=True)

        out_name = output_filename or f"adjusted_{filepath.name}"
        out_path = OUTPUT_DIR / out_name
        img.save(out_path)
        img.close()

        changes = []
        if brightness != 1.0: changes.append(f"яркость={brightness}")
        if contrast != 1.0: changes.append(f"контраст={contrast}")
        if sharpness != 1.0: changes.append(f"резкость={sharpness}")
        if rotate: changes.append(f"поворот={rotate}°")

        return f"✓ {out_name}: {', '.join(changes)}"
    except Exception as e:
        return f"Ошибка: {e}"


@tool
def image_analyze(filename: str, question: str = "Что изображено на этой картинке? Опиши подробно.") -> str:
    """Анализировать содержимое изображения с помощью AI Vision. Отправляет картинку в модель для распознавания.

    Args:
        filename: Имя файла изображения
        question: Вопрос о содержимом изображения
    """
    import base64
    try:
        filepath = _resolve_file(filename)
        if not filepath:
            return f"Файл не найден: {filename}"

        # Определить MIME тип
        ext = filepath.suffix.lower()
        mime_map = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                    ".webp": "image/webp", ".gif": "image/gif"}
        mime = mime_map.get(ext, "image/jpeg")

        # Кодируем в base64
        with open(filepath, "rb") as f:
            img_b64 = base64.b64encode(f.read()).decode("utf-8")

        # Отправляем в текущую модель через OpenAI-совместимый API
        import urllib.request
        api_key = ""
        base_url = ""
        model = ""

        # Читаем настройки из settings.json
        settings_file = BASE_DIR / "settings.json"
        if settings_file.exists():
            try:
                with open(settings_file, "r", encoding="utf-8") as sf:
                    s = json.load(sf)
                api_key = s.get("api_key", "")
                base_url = s.get("base_url", "")
                model = s.get("model", "")
            except Exception:
                pass

        # Fallback на env
        api_key = api_key or os.getenv("PROXYAPI_KEY", "")
        base_url = base_url or os.getenv("BASE_URL", "https://openai.api.proxyapi.ru/v1")
        model = model or os.getenv("VISION_MODEL", "claude-3-5-haiku-20241022")

        # Убираем /v1 если есть — добавим сами
        base_url = base_url.rstrip("/")
        if not base_url.endswith("/v1"):
            base_url += "/v1"

        if not api_key:
            return "⚠️ Нет API-ключа для Vision-анализа"

        logger.info(f"Vision: model={model}, url={base_url}")

        payload = {
            "model": model,
            "max_tokens": 1024,
            "messages": [{
                "role": "user",
                "content": [
                    {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{img_b64}"}},
                    {"type": "text", "text": question}
                ]
            }]
        }

        data = json.dumps(payload).encode("utf-8")
        req = urllib.request.Request(
            f"{base_url}/chat/completions",
            data=data,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Bearer {api_key}"
            }
        )

        with urllib.request.urlopen(req, timeout=30) as resp:
            result = json.loads(resp.read().decode("utf-8"))

        answer = result.get("choices", [{}])[0].get("message", {}).get("content", "")
        return answer or "Не удалось получить описание"

    except Exception as e:
        return f"⚠️ Ошибка анализа: {e}"


@tool
def fetch_url(url: str) -> str:
    """Открыть URL и прочитать содержимое веб-страницы.

    Args:
        url: URL страницы (https://...)
    """
    import urllib.request
    import html
    try:
        if not url.startswith("http"):
            url = "https://" + url

        req = urllib.request.Request(url, headers={
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36"
        })
        with urllib.request.urlopen(req, timeout=15) as resp:
            raw = resp.read(100000).decode("utf-8", errors="replace")

        # Простое извлечение текста из HTML
        import re as _re
        # Убираем script и style
        raw = _re.sub(r'<script[^>]*>.*?</script>', '', raw, flags=_re.DOTALL)
        raw = _re.sub(r'<style[^>]*>.*?</style>', '', raw, flags=_re.DOTALL)
        # Убираем теги
        text = _re.sub(r'<[^>]+>', ' ', raw)
        text = html.unescape(text)
        # Нормализуем пробелы
        text = _re.sub(r'\s+', ' ', text).strip()

        if len(text) > 8000:
            text = text[:8000] + "\n\n... (обрезано)"

        return f"URL: {url}\n\n{text}"
    except Exception as e:
        return f"Ошибка загрузки {url}: {e}"


# ============ BROWSER (Selenium, опционально) ============

_browser_driver = None

def _get_browser():
    """Получить или создать экземпляр браузера."""
    global _browser_driver
    if _browser_driver:
        try:
            _browser_driver.title  # проверяем что жив
            return _browser_driver
        except Exception:
            _browser_driver = None

    if not SELENIUM_AVAILABLE:
        return None

    options = webdriver.ChromeOptions()
    options.add_argument("--headless=new")
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-gpu")
    options.add_argument("--window-size=1280,900")
    options.add_argument("--user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36")

    try:
        _browser_driver = webdriver.Chrome(options=options)
        _browser_driver.set_page_load_timeout(20)
        _browser_driver.implicitly_wait(5)
        return _browser_driver
    except Exception as e:
        logger.error(f"Ошибка запуска Chrome: {e}")
        return None


def _close_browser():
    global _browser_driver
    if _browser_driver:
        try:
            _browser_driver.quit()
        except Exception:
            pass
        _browser_driver = None


@tool
def browser_open(url: str) -> str:
    """Открыть URL в браузере (с поддержкой JavaScript). Используй для сайтов, которые не работают через fetch_url.

    Args:
        url: URL страницы
    """
    if not SELENIUM_AVAILABLE:
        return "⚠️ Selenium не установлен. pip install selenium"

    driver = _get_browser()
    if not driver:
        return "⚠️ Не удалось запустить Chrome. Убедитесь что Chrome/Chromium установлен."

    try:
        if not url.startswith("http"):
            url = "https://" + url
        driver.get(url)
        import time; time.sleep(2)  # дать JS отработать

        # Извлечь текст страницы
        text = driver.find_element(By.TAG_NAME, "body").text
        title = driver.title

        if len(text) > 8000:
            text = text[:8000] + "\n\n... (обрезано)"

        return f"Страница: {title}\nURL: {url}\n\n{text}"
    except TimeoutException:
        return f"⚠️ Таймаут загрузки: {url}"
    except Exception as e:
        return f"⚠️ Ошибка: {e}"


@tool
def browser_click(selector: str, by: str = "css") -> str:
    """Кликнуть на элемент в открытой странице браузера.

    Args:
        selector: CSS-селектор или XPath элемента
        by: Тип селектора — "css", "xpath", "text" (поиск по тексту ссылки)
    """
    if not SELENIUM_AVAILABLE:
        return "⚠️ Selenium не установлен"

    driver = _get_browser()
    if not driver:
        return "⚠️ Браузер не запущен. Сначала используй browser_open"

    try:
        if by == "xpath":
            el = driver.find_element(By.XPATH, selector)
        elif by == "text":
            el = driver.find_element(By.LINK_TEXT, selector)
        else:
            el = driver.find_element(By.CSS_SELECTOR, selector)

        el.click()
        import time; time.sleep(1)

        text = driver.find_element(By.TAG_NAME, "body").text
        if len(text) > 5000:
            text = text[:5000] + "\n..."

        return f"Клик выполнен. Текущая страница: {driver.title}\n\n{text}"
    except NoSuchElementException:
        return f"⚠️ Элемент не найден: {selector}"
    except Exception as e:
        return f"⚠️ Ошибка: {e}"


@tool
def browser_fill(selector: str, value: str, submit: bool = False) -> str:
    """Заполнить поле ввода на странице и опционально отправить форму.

    Args:
        selector: CSS-селектор поля ввода
        value: Текст для ввода
        submit: Нажать Enter после ввода
    """
    if not SELENIUM_AVAILABLE:
        return "⚠️ Selenium не установлен"

    driver = _get_browser()
    if not driver:
        return "⚠️ Браузер не запущен"

    try:
        el = driver.find_element(By.CSS_SELECTOR, selector)
        el.clear()
        el.send_keys(value)
        if submit:
            el.send_keys(Keys.RETURN)
            import time; time.sleep(2)

        text = driver.find_element(By.TAG_NAME, "body").text
        if len(text) > 5000:
            text = text[:5000] + "\n..."

        return f"Поле заполнено. Страница: {driver.title}\n\n{text}"
    except NoSuchElementException:
        return f"⚠️ Поле не найдено: {selector}"
    except Exception as e:
        return f"⚠️ Ошибка: {e}"


@tool
def browser_extract(selector: str) -> str:
    """Извлечь текст из конкретного элемента на странице.

    Args:
        selector: CSS-селектор элемента (например ".price", "#content", "h1")
    """
    if not SELENIUM_AVAILABLE:
        return "⚠️ Selenium не установлен"

    driver = _get_browser()
    if not driver:
        return "⚠️ Браузер не запущен"

    try:
        elements = driver.find_elements(By.CSS_SELECTOR, selector)
        if not elements:
            return f"⚠️ Элементы не найдены: {selector}"

        results = []
        for el in elements[:20]:  # макс 20 элементов
            text = el.text.strip()
            if text:
                results.append(text)

        if not results:
            return f"⚠️ Элементы найдены ({len(elements)}), но пустые"

        return f"Найдено {len(results)} элементов:\n\n" + "\n---\n".join(results)
    except Exception as e:
        return f"⚠️ Ошибка: {e}"


@tool
def browser_screenshot(filename: str = "screenshot.png") -> str:
    """Сделать скриншот текущей страницы в браузере.

    Args:
        filename: Имя файла для сохранения
    """
    if not SELENIUM_AVAILABLE:
        return "⚠️ Selenium не установлен"

    driver = _get_browser()
    if not driver:
        return "⚠️ Браузер не запущен"

    try:
        out_path = OUTPUT_DIR / filename
        driver.save_screenshot(str(out_path))
        return f"Скриншот сохранён: {out_path}"
    except Exception as e:
        return f"⚠️ Ошибка: {e}"


# ============ AGENT ============

ALL_TOOLS = [
    web_search, fetch_url, bash_execute, create_file, view_file, list_files, python_execute,
    # Excel
    excel_create, excel_add_formulas, excel_style,
    excel_read, excel_read_structured, excel_edit_cell, excel_from_csv,
    excel_create_pivot, excel_pivot_analyze,
    # PDF
    pdf_read, pdf_info, pdf_extract_pages,
    # Word
    docx_read, docx_create, docx_to_pdf,
    # Изображения
    image_info, image_resize, image_convert, image_crop, image_adjust, image_analyze,
]

# Добавить браузерные инструменты если Selenium доступен
if SELENIUM_AVAILABLE:
    ALL_TOOLS.extend([browser_open, browser_click, browser_fill, browser_extract, browser_screenshot])
    logger.info(f"🌐 Selenium доступен — добавлено 5 браузерных инструментов")
else:
    logger.info(f"ℹ️ Selenium не установлен — браузерные инструменты отключены")


def create_claude_agent(
    api_key: Optional[str] = None,
    base_url: str = "https://openai.api.proxyapi.ru/v1",
    model: str = "claude-3-5-haiku-20241022",
    use_memory: bool = True,
    temperature: float = 0,
):
    """Создать агента с памятью.

    Args:
        api_key: API ключ (или из переменной окружения PROXYAPI_KEY)
        base_url: URL прокси-сервера
        model: Идентификатор модели
        use_memory: Использовать MemorySaver для сохранения контекста
    """
    api_key = api_key or os.getenv("PROXYAPI_KEY")
    if not api_key:
        raise ValueError("Добавьте PROXYAPI_KEY в .env файл")

    # Модели-рассуждатели не поддерживают temperature
    no_temp_models = ["deepseek-reasoner", "o1", "o1-mini", "o1-preview", "o3", "o3-mini"]
    skip_temp = any(m in model.lower() for m in no_temp_models)

    model_kwargs = {
        "model": model,
        "api_key": api_key,
        "base_url": base_url,
        "max_completion_tokens": 8192,
    }
    if not skip_temp:
        model_kwargs["temperature"] = temperature

    model_instance = ChatOpenAI(**model_kwargs)

    checkpointer = MemorySaver() if use_memory else None

    agent = create_agent(
        model=model_instance, tools=ALL_TOOLS,
        system_prompt=SYSTEM_PROMPT, checkpointer=checkpointer,
    )

    return agent


# ============ HELPERS ============

def make_session_config(thread_id: Optional[str] = None) -> Dict[str, Any]:
    """Создаёт конфиг сессии с уникальным thread_id."""
    tid = thread_id or f"session_{uuid.uuid4().hex[:8]}"
    return {"configurable": {"thread_id": tid}}


def run_agent(
    agent,
    query: str,
    session_config: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    """Запрос к агенту с сохранением контекста."""
    config = session_config or make_session_config()
    logger.info(f"Запрос: {query[:100]}...")
    result = agent.invoke(
        {"messages": [{"role": "user", "content": query}]},
        config=config,
    )
    answer = result["messages"][-1].content
    logger.info(f"Ответ получен ({len(answer)} символов)")
    return {"output": answer, "messages": result["messages"]}


def run_agent_stream(
    agent,
    query: str,
    session_config: Optional[Dict[str, Any]] = None,
    on_token=None,
    stop_event=None,
) -> Dict[str, Any]:
    """Потоковый запрос к агенту. on_token(str) вызывается для каждого чанка."""
    config = session_config or make_session_config()
    logger.info(f"Стрим запрос: {query[:100]}...")

    full_answer = ""
    try:
        for event in agent.stream(
            {"messages": [{"role": "user", "content": query}]},
            config=config,
            stream_mode="messages",
        ):
            if stop_event and stop_event.is_set():
                logger.info("Стрим остановлен пользователем")
                break

            # event может быть (message, metadata) tuple
            msg = event[0] if isinstance(event, tuple) else event
            content = getattr(msg, 'content', '')
            if content and isinstance(content, str):
                full_answer += content
                if on_token:
                    on_token(content)
    except Exception as e:
        logger.error(f"Ошибка стрима: {e}")
        if not full_answer:
            raise

    if not full_answer:
        full_answer = "Нет ответа"

    return {"output": full_answer}


# ============ MAIN ============

if __name__ == "__main__":
    print("🤖 Инициализация агента v3 (с доработками)...\n")

    if not EXCEL_AVAILABLE:
        print("⚠️  Excel библиотеки не найдены")
        print("   pip install openpyxl pandas\n")

    try:
        agent = create_claude_agent(use_memory=True)
        # Одна сессия на весь чат — контекст сохраняется между сообщениями
        session = make_session_config("interactive_chat")

        print("✅ Агент готов!\n")
        print(f"📁 Outputs: {OUTPUT_DIR}")
        print(f"📁 Work:    {WORK_DIR}\n")
        print("Команды:  'quit' — выход  |  'files' — список файлов\n")

        while True:
            try:
                user_input = input("Ты: ").strip()
                if not user_input:
                    continue
                if user_input.lower() in ("quit", "exit", "выход"):
                    print("\n👋 До свидания!")
                    break
                if user_input.lower() == "files":
                    print(list_files.invoke({"directory": "outputs"}))
                    continue

                run_agent(agent, user_input, session_config=session)

            except KeyboardInterrupt:
                print("\n\n👋 До свидания!")
                break
            except Exception as e:
                logger.error(f"Ошибка: {e}", exc_info=True)

    except Exception as e:
        logger.error(f"Не удалось запустить: {e}", exc_info=True)
